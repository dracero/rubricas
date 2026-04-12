"""
Unified Qdrant Service + ADK Tool Functions.

Consolidates the QdrantService that was previously duplicated across
generator, evaluator, and corrector agents into a single module.
Uses Gemini gemini-embedding-001 for embeddings (3072 dimensions).
"""

import os
import hashlib
import logging
from typing import Dict, List, Any, Optional
from datetime import datetime
from collections import defaultdict

from google import genai

try:
    from qdrant_client import QdrantClient
    from qdrant_client.http import models as qmodels
    QDRANT_AVAILABLE = True
except ImportError:
    QdrantClient = None
    QDRANT_AVAILABLE = False

from common.config import ConfiguracionColaba, traceable, get_current_run_tree
from app.domain import (
    Entidad, Relacion, Ontologia,
    parsear_json_con_fallback,
)

logger = logging.getLogger(__name__)

# Gemini embedding config
EMBEDDING_MODEL = "gemini-embedding-001"
EMBEDDING_DIMENSION = 3072
VERSION = "FIX_QDRANT_3072_V1"

# ============================================================================
# Module-level singleton
# ============================================================================
_qdrant_service: Optional["QdrantService"] = None


def _get_qdrant_service() -> "QdrantService":
    """Get or create the module-level QdrantService singleton."""
    global _qdrant_service
    if _qdrant_service is None:
        config = ConfiguracionColaba()
        _qdrant_service = QdrantService(config)
    return _qdrant_service


# ============================================================================
# QDRANT SERVICE
# ============================================================================

class QdrantService:
    """Unified Qdrant vector DB service: embeddings, save, search."""

    def __init__(self, config: ConfiguracionColaba):
        self.config = config
        self.collection_name = "rubricas_entidades"

        # Qdrant client
        if config.QDRANT_URL and QDRANT_AVAILABLE:
            self.client = QdrantClient(
                url=config.QDRANT_URL,
                api_key=config.QDRANT_API_KEY
            )
        else:
            self.client = None
            logger.warning("⚠️ Sin conexión a Qdrant")

        # Gemini embedding client
        self.genai_client = genai.Client(api_key=config.GOOGLE_API_KEY)
        logger.info(f"✅ Gemini embedding model: {EMBEDDING_MODEL} ({EMBEDDING_DIMENSION}d) [VERSION: {VERSION}]")

        self._init_collection()

    def _init_collection(self):
        """Create the Qdrant collection if it doesn't exist, or recreate if dimensions mismatch."""
        if not self.client:
            return
        try:
            collections = self.client.get_collections()
            exists = any(
                c.name == self.collection_name
                for c in collections.collections
            )
            
            if exists:
                # Check dimensions
                info = self.client.get_collection(self.collection_name)
                current_dim = info.config.params.vectors.size
                if current_dim != EMBEDDING_DIMENSION:
                    logger.warning(
                        f"⚠️ Dimension mismatch in {self.collection_name}: "
                        f"current={current_dim}, expected={EMBEDDING_DIMENSION}. "
                        "Recreating collection..."
                    )
                    self.client.delete_collection(self.collection_name)
                    exists = False

            if not exists:
                logger.info(f"📦 Creating Qdrant collection: {self.collection_name} (dim={EMBEDDING_DIMENSION})")
                self.client.create_collection(
                    collection_name=self.collection_name,
                    vectors_config=qmodels.VectorParams(
                        size=EMBEDDING_DIMENSION,
                        distance=qmodels.Distance.COSINE
                    )
                )
        except Exception as e:
            logger.error(f"⚠️ Error initializing Qdrant: {e}")

    def clear_collection(self):
        """Delete and recreate the Qdrant collection."""
        if not self.client:
            logger.warning("⚠️ No Qdrant client — skipping clear")
            return False
        try:
            collections = self.client.get_collections()
            exists = any(
                c.name == self.collection_name
                for c in collections.collections
            )
            if exists:
                logger.info(f"🗑️ Deleting Qdrant collection: {self.collection_name}")
                self.client.delete_collection(collection_name=self.collection_name)

            logger.info(f"📦 Recreating Qdrant collection: {self.collection_name}")
            self.client.create_collection(
                collection_name=self.collection_name,
                vectors_config=qmodels.VectorParams(
                    size=EMBEDDING_DIMENSION,
                    distance=qmodels.Distance.COSINE
                )
            )
            logger.info("✅ Qdrant collection cleared and recreated")
            return True
        except Exception as e:
            logger.error(f"❌ Error clearing Qdrant collection: {e}")
            return False

    def embed(self, text: str) -> List[float]:
        """Generate embedding vector for text using Gemini API."""
        result = self.genai_client.models.embed_content(
            model=EMBEDDING_MODEL,
            contents=text,
        )
        return result.embeddings[0].values

    @traceable(name="QdrantService.save_ontology", run_type="chain")
    def save_ontology(self, ontologia: Ontologia) -> bool:
        """Save ontology entities and relations to Qdrant."""
        if not self.client:
            logger.warning("⚠️ No Qdrant client — skipping save")
            return False

        points = []

        relations_by_entity = defaultdict(list)
        for rel in ontologia.relaciones:
            relations_by_entity[rel.origen].append(rel.to_dict())

        for entidad in ontologia.entidades:
            point_id = hashlib.md5(entidad.nombre.encode()).hexdigest()
            text_for_embedding = f"{entidad.nombre}: {entidad.contexto}"
            vector = self.embed(text_for_embedding)

            payload = entidad.to_dict()
            payload["relaciones_salientes"] = relations_by_entity[entidad.nombre]

            points.append(qmodels.PointStruct(
                id=point_id,
                vector=vector,
                payload=payload
            ))

        if points:
            try:
                run_tree = get_current_run_tree()
                if run_tree:
                    run_tree.extra = run_tree.extra or {}
                    run_tree.extra.update({
                        "qdrant_operation": "upsert",
                        "collection_name": self.collection_name,
                        "num_entities": len(ontologia.entidades),
                        "num_relations": len(ontologia.relaciones),
                        "num_points": len(points),
                        "vector_dimension": EMBEDDING_DIMENSION,
                    })

                self.client.upsert(
                    collection_name=self.collection_name,
                    points=points
                )
                logger.info(f"✅ Saved {len(points)} entities to Qdrant")
                return True
            except Exception as e:
                logger.error(f"❌ Error saving to Qdrant: {e}")
                run_tree = get_current_run_tree()
                if run_tree:
                    run_tree.error = str(e)
                return False
        return False

    @traceable(name="QdrantService.search", run_type="retriever")
    def search(self, query: str, limit: int = 10, score_threshold: float = 0.5) -> List[Dict]:
        """Search for similar entities by vector."""
        if not self.client:
            return []

        vector = self.embed(query)

        try:
            result = self.client.query_points(
                collection_name=self.collection_name,
                query=vector,
                limit=limit,
                score_threshold=score_threshold
            )
            hits = result.points if hasattr(result, 'points') else result

            resultados = []
            for hit in hits:
                payload = hit.payload.copy() if hit.payload else {}
                payload['score'] = hit.score
                resultados.append(payload)

            avg_score = (
                sum(r['score'] for r in resultados) / len(resultados)
                if resultados else 0
            )

            run_tree = get_current_run_tree()
            if run_tree:
                run_tree.extra = run_tree.extra or {}
                run_tree.extra.update({
                    "qdrant_operation": "search",
                    "collection_name": self.collection_name,
                    "query": query[:100],
                    "limit": limit,
                    "score_threshold": score_threshold,
                    "num_results": len(resultados),
                    "avg_score": round(avg_score, 3),
                    "top_scores": [round(r['score'], 3) for r in resultados[:5]],
                })

            logger.info(
                f"📊 Qdrant search: {len(resultados)} hits, "
                f"avg_score: {avg_score:.3f}"
            )
            return resultados
        except Exception as e:
            logger.error(f"⚠️ Qdrant search error: {e}")
            run_tree = get_current_run_tree()
            if run_tree:
                run_tree.error = str(e)
            return []


# ============================================================================
# ADK TOOL FUNCTIONS (shared by all skills)
# ============================================================================

def guardar_ontologia_en_qdrant(ontologia_json: str) -> str:
    """Parses an ontology JSON and saves entities/relations to Qdrant.

    Use this tool after extracting entities and relations from a normative document.
    The JSON must have the structure:
    {
        "entidades": [{"nombre": "...", "tipo": "...", "contexto": "...", "propiedades": {...}}],
        "relaciones": [{"origen": "...", "destino": "...", "tipo": "...", "propiedades": {...}}]
    }

    Args:
        ontologia_json: JSON string with the extracted ontology (entidades + relaciones).

    Returns:
        A confirmation message with the number of entities and relations saved.
    """
    try:
        data = parsear_json_con_fallback(ontologia_json)

        entidades = []
        for e in data.get("entidades", []):
            entidades.append(Entidad(
                nombre=e.get("nombre", "Desconocido"),
                tipo=e.get("tipo", "Desconocido"),
                propiedades=e.get("propiedades", {}),
                contexto=e.get("contexto", ""),
                fecha_creacion=datetime.now().isoformat()
            ))

        relaciones = []
        for r in data.get("relaciones", []):
            relaciones.append(Relacion(
                origen=r.get("origen", "Desconocido"),
                destino=r.get("destino", "Desconocido"),
                tipo=r.get("tipo", "Desconocido"),
                propiedades=r.get("propiedades", {})
            ))

        ontologia = Ontologia(
            entidades=entidades,
            relaciones=relaciones,
            metadata={"source": "skill_ontologo"}
        )

        qdrant = _get_qdrant_service()
        
        # Clear existing collection before saving new ontology
        logger.info("🗑️ Limpiando colección Qdrant antes de guardar nueva ontología...")
        qdrant.clear_collection()
        
        success = qdrant.save_ontology(ontologia)

        if success:
            return (
                f"✅ Ontología guardada exitosamente en Qdrant: "
                f"{len(entidades)} entidades, {len(relaciones)} relaciones."
            )
        else:
            return "⚠️ No se pudieron guardar entidades en Qdrant."

    except Exception as e:
        logger.error(f"Error in guardar_ontologia_en_qdrant: {e}")
        return f"❌ Error guardando ontología: {str(e)}"


def buscar_contexto_qdrant(consulta: str) -> str:
    """Searches the Qdrant vector database for normative context relevant to the query.

    Use this tool to retrieve knowledge from previously indexed normative documents.
    Returns entities with their similarity scores, descriptions, and relationships.

    Args:
        consulta: The search query describing what normative context is needed.

    Returns:
        A formatted string with the relevant normative context found in Qdrant.
    """
    try:
        qdrant = _get_qdrant_service()
        results = qdrant.search(consulta, limit=10, score_threshold=0.4)

        if not results:
            return "No se encontró contexto normativo relevante en la base de conocimiento."

        lines = [f"📚 Contexto normativo encontrado ({len(results)} documentos):\n"]
        for item in results:
            score = item.get('score', 0)
            nombre = item.get('nombre', 'N/A')
            contexto = item.get('contexto', '')[:300]
            lines.append(f"- [{score:.3f}] **{nombre}**: {contexto}")

            for rel in item.get('relaciones_salientes', []):
                lines.append(
                    f"  → {rel.get('tipo', '?')} → {rel.get('destino', '?')}"
                )

        return "\n".join(lines)

    except Exception as e:
        logger.error(f"Error in buscar_contexto_qdrant: {e}")
        return f"❌ Error buscando contexto: {str(e)}"

def leer_rubrica_subida(rubric_id: str) -> str:
    """Reads the content of a rubric file that was previously uploaded via the frontend.

    Use this tool when you need to access the text of a rubric that the user
    uploaded through the UI. The rubric_id is provided by the frontend after upload.
    Supports DOCX, PDF, TXT, and MD formats.

    Args:
        rubric_id: The unique ID of the uploaded rubric (provided after upload).

    Returns:
        The text content of the rubric file, or an error message if not found.
    """
    from pathlib import Path

    upload_dir = Path("/tmp/rubricas_uploads")

    # Try common extensions (including DOCX and PDF)
    for ext in [".docx", ".pdf", ".txt", ".md"]:
        rubric_path = upload_dir / f"rubric_{rubric_id}{ext}"
        if rubric_path.exists():
            try:
                if ext == ".docx":
                    from docx import Document
                    doc = Document(str(rubric_path))
                    text_parts = []
                    for para in doc.paragraphs:
                        if para.text.strip():
                            text_parts.append(para.text)
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = " | ".join(cell.text.strip() for cell in row.cells)
                            text_parts.append(row_text)
                    text = "\n".join(text_parts)
                    if not text.strip():
                        return "⚠️ El archivo DOCX de la rúbrica no contiene texto."
                elif ext == ".pdf":
                    import pypdf
                    reader = pypdf.PdfReader(str(rubric_path))
                    text_parts = []
                    for page in reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                    text = "\n".join(text_parts)
                    if not text.strip():
                        return "⚠️ El PDF de la rúbrica no contiene texto extraíble."
                else:
                    text = rubric_path.read_text(encoding="utf-8")
                
                logger.info(f"📖 Rubric read: rubric_{rubric_id}{ext} ({len(text)} chars)")
                return text
            except ImportError as ie:
                return f"❌ Dependencia faltante para leer {ext}: {str(ie)}"
            except Exception as e:
                return f"❌ Error leyendo rúbrica: {str(e)}"

    return f"❌ Rúbrica no encontrada con ID: {rubric_id}"


def leer_documento_subido(document_id: str) -> str:
    """Reads and extracts text from a PDF document that was previously uploaded via the frontend.

    Use this tool when you need to access the content of a document that the user
    uploaded for evaluation or analysis. The document_id is provided by the frontend after upload.
    Supports PDF files — text is automatically extracted from all pages.

    Args:
        document_id: The unique ID of the uploaded document (provided after upload).

    Returns:
        The extracted text content of the document, or an error message if not found.
    """
    from pathlib import Path

    upload_dir = Path("/tmp/rubricas_uploads")

    # Check both naming patterns: {id}.pdf and doc_{id}.pdf
    for pattern in [f"doc_{document_id}.pdf", f"{document_id}.pdf"]:
        doc_path = upload_dir / pattern
        if doc_path.exists():
            try:
                import pypdf
                reader = pypdf.PdfReader(str(doc_path))
                text_parts = []
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                full_text = "\n".join(text_parts)
                logger.info(f"📖 Document read: {pattern} ({len(full_text)} chars)")
                return full_text if full_text.strip() else "⚠️ El PDF no contiene texto extraíble."
            except ImportError:
                return "❌ pypdf no está instalado. No se puede leer el PDF."
            except Exception as e:
                return f"❌ Error leyendo documento: {str(e)}"

    return f"❌ Documento no encontrado con ID: {document_id}"


# Registry of all available tool functions (for skill_loader)
TOOL_REGISTRY: Dict[str, Any] = {
    "guardar_ontologia_en_qdrant": guardar_ontologia_en_qdrant,
    "buscar_contexto_qdrant": buscar_contexto_qdrant,
    "leer_rubrica_subida": leer_rubrica_subida,
    "leer_documento_subido": leer_documento_subido,
}
