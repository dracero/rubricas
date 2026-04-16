"""
DOCX Converter module for RubricAI.

Provides utilities for:
- Converting Markdown text to DOCX documents (md_to_docx)
- Extracting text from DOCX files (extract_text_from_docx)
- Detecting rubric tables in agent responses (detect_rubric_in_response)
"""

import logging
import re
from html.parser import HTMLParser
from typing import List, Optional, Tuple
from zipfile import BadZipFile

import markdown
from docx import Document
from docx.shared import Pt, RGBColor, Mm
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Internal HTML → DOCX parser
# ---------------------------------------------------------------------------

class _DocxHTMLParser(HTMLParser):
    """Custom HTMLParser that builds a python-docx Document from HTML tags.

    Maintains state while traversing the HTML produced by the ``markdown``
    library and creates the corresponding DOCX elements (headings, paragraphs,
    tables, bold/italic runs, list items).
    """

    def __init__(self, document: Document) -> None:
        super().__init__()
        self.doc = document

        # Current paragraph / run state
        self._current_paragraph = None
        self._bold = False
        self._italic = False

        # Tag stack for nesting awareness
        self._tag_stack: List[str] = []

        # Table state
        self._in_table = False
        self._in_thead = False
        self._in_tbody = False
        self._current_row: List[str] = []
        self._table_rows: List[List[str]] = []
        self._is_header_row = False
        self._current_cell_text = ""
        self._in_cell = False

    # -- helpers ----------------------------------------------------------

    def _flush_text(self, text: str) -> None:
        """Add *text* as a run on the current paragraph, applying formatting."""
        if self._current_paragraph is None:
            self._current_paragraph = self.doc.add_paragraph()
        run = self._current_paragraph.add_run(text)
        if self._bold:
            run.bold = True
        if self._italic:
            run.italic = True

    def _finish_table(self) -> None:
        """Render the accumulated table rows into a Word table."""
        if not self._table_rows:
            return

        num_cols = max(len(r) for r in self._table_rows)
        # Normalise rows to same column count
        for row in self._table_rows:
            while len(row) < num_cols:
                row.append("")

        num_rows = len(self._table_rows)
        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Allow table to auto-fit to page width
        table.autofit = True

        # Apply borders via XML
        tbl = table._tbl
        tbl_pr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
        borders = tbl_pr.makeelement(qn("w:tblBorders"), {})
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            element = borders.makeelement(
                qn(f"w:{edge}"),
                {
                    qn("w:val"): "single",
                    qn("w:sz"): "4",
                    qn("w:space"): "0",
                    qn("w:color"): "AAAAAA",
                },
            )
            borders.append(element)
        tbl_pr.append(borders)

        for row_idx, row_data in enumerate(self._table_rows):
            for col_idx, cell_text in enumerate(row_data):
                cell = table.cell(row_idx, col_idx)
                cell.text = cell_text.strip()
                # Apply smaller font for table cells to fit landscape
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)
                # Style header row (first row): bold + light gray background
                if row_idx == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.size = Pt(8)
                    shading = cell._element.makeelement(
                        qn("w:shd"),
                        {
                            qn("w:val"): "clear",
                            qn("w:color"): "auto",
                            qn("w:fill"): "D9D9D9",
                        },
                    )
                    tc_pr = cell._element.get_or_add_tcPr()
                    tc_pr.append(shading)

        # Reset table state
        self._table_rows = []
        self._current_row = []

    # -- HTMLParser overrides ---------------------------------------------

    def handle_starttag(self, tag: str, attrs: List[Tuple[str, Optional[str]]]) -> None:
        tag = tag.lower()
        self._tag_stack.append(tag)

        if tag in ("h1", "h2", "h3"):
            level = int(tag[1])
            self._current_paragraph = self.doc.add_heading(level=level)
        elif tag == "p":
            if not self._in_table:
                self._current_paragraph = self.doc.add_paragraph()
        elif tag in ("strong", "b"):
            self._bold = True
        elif tag in ("em", "i"):
            self._italic = True
        elif tag == "li":
            self._current_paragraph = self.doc.add_paragraph()
            self._flush_text("• ")
        elif tag == "table":
            self._in_table = True
            self._table_rows = []
            self._current_row = []
        elif tag == "thead":
            self._in_thead = True
        elif tag == "tbody":
            self._in_tbody = True
        elif tag in ("tr",):
            self._current_row = []
            self._is_header_row = self._in_thead
        elif tag in ("th", "td"):
            self._in_cell = True
            self._current_cell_text = ""

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()

        if tag in ("strong", "b"):
            self._bold = False
        elif tag in ("em", "i"):
            self._italic = False
        elif tag in ("h1", "h2", "h3", "p", "li"):
            self._current_paragraph = None
        elif tag in ("th", "td"):
            self._in_cell = False
            self._current_row.append(self._current_cell_text)
            self._current_cell_text = ""
        elif tag == "tr":
            if self._current_row:
                self._table_rows.append(self._current_row)
            self._current_row = []
        elif tag == "thead":
            self._in_thead = False
        elif tag == "tbody":
            self._in_tbody = False
        elif tag == "table":
            self._finish_table()
            self._in_table = False

        # Pop tag stack
        if self._tag_stack and self._tag_stack[-1] == tag:
            self._tag_stack.pop()

    def handle_data(self, data: str) -> None:
        if self._in_cell:
            self._current_cell_text += data
        elif self._current_paragraph is not None:
            self._flush_text(data)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def md_to_docx(md_text: str, output_path: str) -> None:
    """Convert Markdown text to a landscape-oriented DOCX file.

    Parses Markdown into HTML using the ``markdown`` library (with *tables*
    and *fenced_code* extensions), then walks the HTML with a custom
    :class:`HTMLParser` to build a Word document via ``python-docx``.

    The document is set to **landscape** orientation (A4) so that wide
    evaluation matrices fit comfortably.

    Handles:
    - Headings (H1-H3 → Heading 1-3 styles)
    - Tables (with borders and styled header row)
    - Bold / italic formatting
    - Paragraphs and list items

    Parameters
    ----------
    md_text : str
        The Markdown source text.
    output_path : str
        Filesystem path where the ``.docx`` file will be written.
    """
    doc = Document()

    # Set landscape orientation (A4: 210 × 297 mm → swap width/height)
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    # A4 dimensions swapped for landscape
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    # Tighter margins for more table space
    section.left_margin = Mm(15)
    section.right_margin = Mm(15)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)

    if md_text and md_text.strip():
        html_content = markdown.markdown(
            md_text, extensions=["tables", "fenced_code"]
        )
        parser = _DocxHTMLParser(doc)
        parser.feed(html_content)
        logger.info("Converted Markdown (%d chars) to landscape DOCX: %s", len(md_text), output_path)
    else:
        logger.info("Empty Markdown input — writing empty DOCX: %s", output_path)

    doc.save(output_path)


def extract_text_from_docx(docx_path: str) -> str:
    """Extract all text from a DOCX file.

    Iterates over every paragraph and every table cell in the document,
    concatenating their text content separated by newlines.

    Parameters
    ----------
    docx_path : str
        Filesystem path to the ``.docx`` file.

    Returns
    -------
    str
        The extracted text.  Returns an empty string for empty documents.

    Raises
    ------
    fastapi.HTTPException
        If the file is not a valid ZIP / DOCX (``BadZipFile``).
    """
    try:
        doc = Document(docx_path)
    except BadZipFile:
        logger.error("Corrupted or invalid DOCX file: %s", docx_path)
        # Import here to avoid hard dependency on FastAPI when used standalone
        from fastapi import HTTPException
        raise HTTPException(
            status_code=400,
            detail="El archivo DOCX está corrupto o no es un archivo DOCX válido.",
        )
    except Exception as e:
        logger.error("Error opening DOCX file %s: %s", docx_path, e)
        from fastapi import HTTPException
        raise HTTPException(
            status_code=500,
            detail=f"Error al leer el archivo DOCX: {str(e)}",
        )

    parts: List[str] = []

    # Paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts.append(text)

    return "\n".join(parts)


def detect_rubric_in_response(response_text: str) -> bool:
    """Detect whether *response_text* contains a Markdown table.

    The heuristic looks for two patterns that together indicate a
    Markdown table:

    1. At least one line with **two or more** ``|`` characters (column
       separators).
    2. At least one line that looks like a header separator, e.g.
       ``|---|---|``.

    Parameters
    ----------
    response_text : str
        The plain-text (Markdown) response from the agent.

    Returns
    -------
    bool
        ``True`` if both patterns are found, ``False`` otherwise.
    """
    if not response_text:
        return False

    has_pipe_columns = False
    has_header_separator = False

    # Pattern for header separator rows like |---|---| or | --- | --- |
    separator_pattern = re.compile(r"^\s*\|[\s\-:]+(\|[\s\-:]+)+\|?\s*$")

    for line in response_text.splitlines():
        # Check for column separators (at least 2 pipes on a line)
        if line.count("|") >= 2:
            has_pipe_columns = True

        # Check for header separator
        if separator_pattern.match(line):
            has_header_separator = True

        # Short-circuit once both found
        if has_pipe_columns and has_header_separator:
            return True

    return False
